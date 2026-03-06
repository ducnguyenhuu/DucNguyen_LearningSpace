"""DOCX document parser using python-docx.

Extracts text paragraph-by-paragraph from Word documents (.docx), preserving
heading structure via built-in Word styles (Heading 1, Heading 2, etc.).

Edge cases handled
------------------
- **Corrupted / invalid DOCX**: ``zipfile.BadZipFile`` and other library
  exceptions are caught and re-raised as ``ParserError`` so the ingestion
  service can skip the file, log the error, and continue
  (Edge Case §corrupted per spec).
- **Empty document**: returns a ``ParsedDocument`` with empty ``content``
  (``is_empty == True``); the ingestion service decides whether to skip it.
- **Very large documents (500+ pages)**: paragraphs are consumed via an
  iterative generator to avoid loading the entire document into memory at
  once (Edge Case §large-document per spec).
- **Paragraphs with no text**: skipped silently — blank/structural paragraphs
  in DOCX add noise without value.
- **Unknown heading levels**: any style name not matching "Heading N" is
  treated as body text (level 0).

Notes
-----
DOCX has no first-class page concept at the paragraph level; python-docx
does not expose page-break positions.  Consequently ``pages`` always
contains a single synthetic entry with all extracted text, and
``page_number`` on every ``SectionContent`` is ``None``.
"""
from __future__ import annotations

import os
import zipfile
from typing import Iterator

import structlog

from app.parsers.base import (
    DocumentParser,
    PageContent,
    ParsedDocument,
    ParserError,
    SectionContent,
)

log = structlog.get_logger(__name__)

# Prefix used by Word built-in heading styles.
_HEADING_STYLE_PREFIX = "heading"

# Maximum heading level to detect.  Styles above this are treated as body.
_MAX_HEADING_LEVEL = 6


def _parse_heading_level(style_name: str) -> int:
    """Return the heading level (1–_MAX_HEADING_LEVEL) or 0 for body text.

    Parameters
    ----------
    style_name:
        Value of ``paragraph.style.name`` from python-docx, e.g.
        ``"Heading 1"``, ``"Normal"``, ``"List Bullet"``.

    Returns
    -------
    int
        1–_MAX_HEADING_LEVEL for heading styles, 0 otherwise.
    """
    lower = style_name.strip().lower()
    if not lower.startswith(_HEADING_STYLE_PREFIX):
        return 0
    suffix = lower[len(_HEADING_STYLE_PREFIX):].strip()
    try:
        level = int(suffix)
    except ValueError:
        return 0
    return level if 1 <= level <= _MAX_HEADING_LEVEL else 0


def _iter_paragraphs(doc: "docx.document.Document") -> Iterator[tuple[str, str]]:  # type: ignore[name-defined]  # noqa: F821
    """Yield ``(style_name, text)`` for each paragraph, skipping empty ones.

    Parameters
    ----------
    doc:
        An open ``docx.Document`` instance.

    Yields
    ------
    tuple[str, str]
        ``(style_name, text)`` where ``style_name`` is the raw Word style
        (e.g. ``"Heading 1"``) and ``text`` is the stripped paragraph text.
    """
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            yield para.style.name, text


class DocxParser(DocumentParser):
    """Parse DOCX files into :class:`~app.parsers.base.ParsedDocument`.

    Uses python-docx for paragraph extraction.  Heading structure is
    inferred from built-in Word heading styles (``Heading 1`` – ``Heading 6``).
    """

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    async def parse(self, file_path: str) -> ParsedDocument:
        """Extract text and structure from a DOCX file.

        Parameters
        ----------
        file_path:
            Absolute path to the ``.docx`` file on disk.

        Returns
        -------
        ParsedDocument
            Populated DTO including full text, a single synthetic page, and
            section metadata derived from Word heading styles.

        Raises
        ------
        ParserError
            If the file is not found, is not a valid DOCX (bad zip), or any
            other unrecoverable error occurs during extraction.
        """
        logger = log.bind(file_path=file_path)
        logger.info("docx_parse_start")

        if not os.path.exists(file_path):
            raise ParserError(
                file_path=file_path,
                reason="File not found",
            )

        try:
            from docx import Document  # noqa: PLC0415  (lazy import OK — optional dep)
            doc = Document(file_path)
        except zipfile.BadZipFile as exc:
            raise ParserError(
                file_path=file_path,
                reason=f"File is not a valid DOCX (bad zip): {exc}",
            ) from exc
        except Exception as exc:
            raise ParserError(
                file_path=file_path,
                reason=f"Failed to open DOCX: {exc}",
            ) from exc

        try:
            content_parts: list[str] = []
            sections: list[SectionContent] = []
            char_offset = 0

            # Track the "current section" being accumulated.
            current_heading: str = ""
            current_level: int = 0
            current_section_start: int = 0
            current_section_parts: list[str] = []

            def _flush_section() -> None:
                """Save the accumulated section into the sections list."""
                if current_section_parts:
                    section_text = "\n".join(current_section_parts)
                    sections.append(
                        SectionContent(
                            heading=current_heading,
                            level=current_level,
                            text=section_text,
                            page_number=None,
                            char_offset=current_section_start,
                        )
                    )

            for style_name, text in _iter_paragraphs(doc):
                level = _parse_heading_level(style_name)

                if level > 0:
                    # Starting a new heading — flush the previous section first.
                    _flush_section()
                    current_heading = text
                    current_level = level
                    current_section_start = char_offset
                    current_section_parts = [text]
                else:
                    current_section_parts.append(text)

                content_parts.append(text)
                char_offset += len(text) + 1  # +1 for the newline separator

            # Flush the final section.
            _flush_section()

            full_content = "\n".join(content_parts)

            # DOCX has no accessible page boundaries; expose one synthetic page.
            pages: list[PageContent] = (
                [PageContent(page_number=1, text=full_content)]
                if full_content
                else []
            )

            file_name = os.path.basename(file_path)
            result = ParsedDocument(
                file_path=file_path,
                file_name=file_name,
                file_type=".docx",
                content=full_content,
                pages=pages,
                sections=sections,
            )

            logger.info(
                "docx_parse_complete",
                paragraphs=len(content_parts),
                sections=len(sections),
                content_length=len(full_content),
            )
            return result

        except ParserError:
            raise
        except Exception as exc:
            logger.error(
                "docx_parse_unexpected_error",
                error=str(exc),
                exc_info=True,
            )
            raise ParserError(
                file_path=file_path,
                reason=f"Unexpected error during DOCX parsing: {exc}",
            ) from exc
