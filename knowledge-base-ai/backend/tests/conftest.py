"""Shared pytest fixtures for all backend tests.

Provides
--------
- Temporary file builders for every supported document format (PDF, DOCX, MD,
  XLSX) so parser unit tests can use real files without external test assets.
- Helpers for corrupt / edge-case file creation.

All fixtures operate under pytest's built-in ``tmp_path`` fixture so files are
automatically cleaned up after each test.
"""
from __future__ import annotations

import os
import textwrap
import zipfile
from io import BytesIO
from typing import Generator

import pytest


# ---------------------------------------------------------------------------
# Markdown fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def md_simple(tmp_path: "os.PathLike[str]") -> str:
    """A minimal Markdown file with ATX headings and body text."""
    content = textwrap.dedent("""\
        # Introduction

        This is the introduction paragraph.

        ## Background

        Some background information here.

        ### Detail

        A deeper detail section.

        This is body text with no heading.
    """)
    p = tmp_path / "simple.md"
    p.write_text(content, encoding="utf-8")
    return str(p)


@pytest.fixture()
def md_setext(tmp_path: "os.PathLike[str]") -> str:
    """A Markdown file using Setext-style headings."""
    content = textwrap.dedent("""\
        Main Title
        ==========

        Some introductory text.

        Section Two
        -----------

        Content under section two.
    """)
    p = tmp_path / "setext.md"
    p.write_text(content, encoding="utf-8")
    return str(p)


@pytest.fixture()
def md_front_matter(tmp_path: "os.PathLike[str]") -> str:
    """A Markdown file with a YAML front-matter block."""
    content = textwrap.dedent("""\
        ---
        title: My Document
        author: Test Author
        date: 2026-03-04
        ---

        # Real Content

        This text should appear in the parsed document.
    """)
    p = tmp_path / "front_matter.md"
    p.write_text(content, encoding="utf-8")
    return str(p)


@pytest.fixture()
def md_empty(tmp_path: "os.PathLike[str]") -> str:
    """An empty Markdown file."""
    p = tmp_path / "empty.md"
    p.write_text("", encoding="utf-8")
    return str(p)


@pytest.fixture()
def md_malformed(tmp_path: "os.PathLike[str]") -> str:
    """Markdown with unusual / malformed content — parser must not raise."""
    content = textwrap.dedent("""\
        ########## Too many hashes (not a heading)
        Normal paragraph text.
        #NoSpaceHeading
        Another paragraph.
        ---
        Not a setext heading because this line is the first.
    """)
    p = tmp_path / "malformed.md"
    p.write_text(content, encoding="utf-8")
    return str(p)


@pytest.fixture()
def md_no_headings(tmp_path: "os.PathLike[str]") -> str:
    """Markdown with only body text and no headings."""
    content = "Just plain text.\nNo headings at all.\nAnother line.\n"
    p = tmp_path / "no_headings.md"
    p.write_text(content, encoding="utf-8")
    return str(p)


# ---------------------------------------------------------------------------
# DOCX fixtures
# ---------------------------------------------------------------------------


def _build_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Build a DOCX in memory using python-docx so heading styles resolve correctly.

    Parameters
    ----------
    paragraphs:
        List of ``(style_name, text)`` tuples.  ``style_name`` should be a
        Word built-in style name such as ``"Normal"`` or ``"Heading 1"``.

    Returns
    -------
    bytes
        Raw DOCX bytes.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    for style_name, text in paragraphs:
        if style_name.lower().startswith("heading"):
            suffix = style_name.strip().split()[-1]
            try:
                level = int(suffix)
            except ValueError:
                level = 1
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text, style="Normal")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture()
def docx_with_headings(tmp_path: "os.PathLike[str]") -> str:
    """A DOCX file with Heading 1, Heading 2, and Normal paragraphs."""
    paragraphs = [
        ("Heading 1", "Introduction"),
        ("Normal", "This is the introduction paragraph."),
        ("Normal", "Second paragraph of the introduction."),
        ("Heading 2", "Background"),
        ("Normal", "Background information here."),
        ("Heading 1", "Conclusion"),
        ("Normal", "Conclusion text."),
    ]
    p = tmp_path / "with_headings.docx"
    p.write_bytes(_build_docx(paragraphs))
    return str(p)


@pytest.fixture()
def docx_no_headings(tmp_path: "os.PathLike[str]") -> str:
    """A DOCX file with only Normal paragraphs (no headings)."""
    paragraphs = [
        ("Normal", "First paragraph - this is plain body text."),
        ("Normal", "Second paragraph."),
        ("Normal", "Third paragraph."),
    ]
    p = tmp_path / "no_headings.docx"
    p.write_bytes(_build_docx(paragraphs))
    return str(p)


@pytest.fixture()
def docx_empty(tmp_path: "os.PathLike[str]") -> str:
    """A DOCX file with no paragraphs (empty document)."""
    p = tmp_path / "empty.docx"
    p.write_bytes(_build_docx([]))
    return str(p)


@pytest.fixture()
def docx_corrupt(tmp_path: "os.PathLike[str]") -> str:
    """A file with a .docx extension but corrupt (non-ZIP) content."""
    p = tmp_path / "corrupt.docx"
    p.write_bytes(b"This is not a valid ZIP file at all!!")
    return str(p)


# ---------------------------------------------------------------------------
# Excel fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def xlsx_multi_sheet(tmp_path: "os.PathLike[str]") -> str:
    """An XLSX workbook with two sheets: one with headers, one without."""
    try:
        from openpyxl import Workbook
    except ImportError:
        pytest.skip("openpyxl not installed")

    wb = Workbook()

    # Sheet 1: has a header row
    ws1 = wb.active
    assert ws1 is not None
    ws1.title = "Inventory"
    ws1.append(["Item", "Quantity", "Price"])
    ws1.append(["Widget A", "10", "5.99"])
    ws1.append(["Widget B", "3", "12.50"])

    # Sheet 2: no header row (all numbers)
    ws2 = wb.create_sheet("Raw Data")
    ws2.append(["100", "200", "300"])
    ws2.append(["400", "500", "600"])

    path = tmp_path / "multi_sheet.xlsx"
    wb.save(str(path))
    return str(path)


@pytest.fixture()
def xlsx_empty_sheet(tmp_path: "os.PathLike[str]") -> str:
    """An XLSX workbook where all sheets are empty."""
    try:
        from openpyxl import Workbook
    except ImportError:
        pytest.skip("openpyxl not installed")

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "EmptySheet"
    # No rows added — sheet is empty.

    path = tmp_path / "empty_sheet.xlsx"
    wb.save(str(path))
    return str(path)


@pytest.fixture()
def xlsx_single_sheet(tmp_path: "os.PathLike[str]") -> str:
    """An XLSX workbook with a single sheet with header + data rows."""
    try:
        from openpyxl import Workbook
    except ImportError:
        pytest.skip("openpyxl not installed")

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Employees"
    ws.append(["Name", "Department", "Salary"])
    ws.append(["Alice", "Engineering", "90000"])
    ws.append(["Bob", "Marketing", "75000"])

    path = tmp_path / "single_sheet.xlsx"
    wb.save(str(path))
    return str(path)


@pytest.fixture()
def xlsx_corrupt(tmp_path: "os.PathLike[str]") -> str:
    """A file with a .xlsx extension but corrupt (non-ZIP) content."""
    p = tmp_path / "corrupt.xlsx"
    p.write_bytes(b"Not a valid zip/xlsx file!")
    return str(p)


@pytest.fixture()
def xls_file(tmp_path: "os.PathLike[str]") -> str:
    """A file with a .xls extension (legacy format — not supported)."""
    p = tmp_path / "legacy.xls"
    p.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")  # OLE2 magic bytes
    return str(p)


# ---------------------------------------------------------------------------
# PDF fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def pdf_single_page(tmp_path: "os.PathLike[str]") -> str:
    """A single-page PDF with extractable text."""
    try:
        import fitz
    except ImportError:
        pytest.skip("PyMuPDF not installed")

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Hello from page one.\nThis is the first paragraph.")
    path = tmp_path / "single_page.pdf"
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture()
def pdf_multi_page(tmp_path: "os.PathLike[str]") -> str:
    """A three-page PDF each containing distinct text."""
    try:
        import fitz
    except ImportError:
        pytest.skip("PyMuPDF not installed")

    doc = fitz.open()
    for i in range(1, 4):
        page = doc.new_page()
        page.insert_text((72, 100), f"Content of page {i}.\nSome text on page {i}.")
    path = tmp_path / "multi_page.pdf"
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture()
def pdf_corrupt(tmp_path: "os.PathLike[str]") -> str:
    """A file with a .pdf extension but containing garbage bytes."""
    p = tmp_path / "corrupt.pdf"
    p.write_bytes(b"%PDF-1.4\nThis is not a real PDF body\x00\xff\xfe")
    return str(p)


# ---------------------------------------------------------------------------
# Shared "file not found" path fixture
# ---------------------------------------------------------------------------


@pytest.fixture()
def nonexistent_file(tmp_path: "os.PathLike[str]") -> str:
    """A path that does not exist on disk."""
    return str(tmp_path / "does_not_exist.pdf")
